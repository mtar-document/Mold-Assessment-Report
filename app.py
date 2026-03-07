import streamlit as st
import re
from datetime import datetime, date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import fitz  # PyMuPDF for PDF reading and merging
from PIL import Image
import os

# Page config
st.set_page_config(
    page_title="Mold Assessment Report Generator - Mold Testing and Removal",
    page_icon="🦠",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #0077B6;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.3rem;
        color: #0077B6;
        border-bottom: 2px solid #0077B6;
        padding-bottom: 0.5rem;
        margin-top: 1.5rem;
    }
    .stButton>button {
        background-color: #0077B6;
        color: white;
        font-size: 1.2rem;
        padding: 0.5rem 2rem;
        border-radius: 5px;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeeba;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown('<h1 class="main-header">🦠 Mold Assessment Report Generator 🦠</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Mold Testing and Removal</p>', unsafe_allow_html=True)

# Initialize session state
if 'lab_results' not in st.session_state:
    st.session_state.lab_results = None

def parse_prolab_pdf(pdf_file):
    """Parse PRO-LAB PDF and extract key information"""
    results = {
        'report_number': '',
        'outdoor_control': {
            'total_spores': 0,
            'penicillium_aspergillus': 0,
            'cladosporium': 0,
            'ascospores': 0,
            'stachybotrys': False,
            'chaetomium': False
        },
        'indoor_samples': [],
        'surface_samples': [],
        'conclusion': 'NOT ELEVATED',
        'mold_present': False,
        'dangerous_mold': False
    }
    
    try:
        pdf_bytes = pdf_file.read()
        pdf_file.seek(0)  # Reset for later use
        
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        doc.close()
        
        # Extract report number
        report_match = re.search(r'Report Number:\s*(\d+)', full_text)
        if report_match:
            results['report_number'] = report_match.group(1)
        
        # Check for ELEVATED or UNUSUAL
        if 'ELEVATED' in full_text.upper():
            results['conclusion'] = 'ELEVATED'
            results['mold_present'] = True
        
        if 'UNUSUAL' in full_text.upper():
            results['mold_present'] = True
        
        # Check for dangerous molds
        if 'Stachybotrys' in full_text or 'STACHYBOTRYS' in full_text:
            results['dangerous_mold'] = True
            results['mold_present'] = True
        
        if 'Chaetomium' in full_text or 'CHAETOMIUM' in full_text:
            results['dangerous_mold'] = True
            results['mold_present'] = True
        
        # Try to extract spore counts using regex
        # Look for Penicillium/Aspergillus counts
        pen_asp_matches = re.findall(r'Penicillium/Aspergillus\s+(\d+)\s+(\d+)', full_text)
        if pen_asp_matches:
            for match in pen_asp_matches:
                count = int(match[1])
                if results['outdoor_control']['penicillium_aspergillus'] == 0:
                    results['outdoor_control']['penicillium_aspergillus'] = count
        
        # Look for total spore counts
        total_matches = re.findall(r'TOTAL SPORES\s+(\d+)\s+(\d+)', full_text)
        if total_matches:
            for match in total_matches:
                count = int(match[1])
                if results['outdoor_control']['total_spores'] == 0:
                    results['outdoor_control']['total_spores'] = count
        
        # Determine if indoor > outdoor (simplified check)
        lines = full_text.split('\n')
        for i, line in enumerate(lines):
            if 'NOT ELEVATED' in line:
                pass  # Keep default
            elif 'ELEVATED' in line and 'NOT' not in line:
                results['conclusion'] = 'ELEVATED'
                results['mold_present'] = True
                
    except Exception as e:
        st.warning(f"Could not fully parse PDF: {str(e)}. Please verify results manually.")
    
    return results

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def make_tight(para):
    """Removes all padding and forced spacing from a paragraph."""
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    return para

def make_top_tight(para)
    """Removes ONLY BEFORE spacing from a paragraph."""
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    return para

def create_report(data, photos, lab_pdf_bytes, lab_results):
    """Generate the Word document report"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # ===== PAGE 1: COVER PAGE =====
    # Company Header  
    contact = make_tight(doc.add_paragraph())
    contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact.add_run("Mold Testing and Removal\n")
    contact.add_run("2031 John West Rd. #119\n")
    contact.add_run("Dallas, TX 75228\n")
    contact.add_run("(817) 718-5086\n")
    contact.add_run("help@moldtestingandremoval.com")
    
    # Title
    make_tight(doc.add_paragraph())
    title = make_tight(doc.add_paragraph())
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MOLD ASSESSMENT REPORT")
    run.font.name = 'Bebas Neue'
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(24, 64, 88)
    
    make_tight(doc.add_paragraph())
    
    # Property Photo
    if photos.get('property'):
        try:
            p = make_tight(doc.add_paragraph())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(photos['property'], width=Inches(5))
        except:
            pass
    
    make_tight(doc.add_paragraph())
    
    # Client Info
    info = make_tight(doc.add_paragraph())
    run = info.add_run("Client & Property:\n")
    run.font.name = 'Bebas Neue'
    run.bold = True
    run.font.color.rgb = RGBColor(24, 64, 88)
    run.font.size = Pt(15)
    info.add_run(f"{data['client_name']}\n")
    info.add_run(f"{data['address']}\n")
    info.add_run(f"{data['city']}, {data['state']} {data['zip']}\n")
    
    info2 = make_tight(doc.add_paragraph())
    run = info2.add_run("Assessment Date: ")
    run.font.name = 'Bebas Neue'
    run.bold = True
    run.font.color.rgb = RGBColor(24, 64, 88)
    run.font.size = Pt(15)
    info2.add_run(data['inspection_date'].strftime("%B %d, %Y"))
    
    info3 = make_tight(doc.add_paragraph())
    run = info3.add_run("Report Date: ")
    run.font.name = 'Bebas Neue'
    run.bold = True
    run.font.color.rgb = RGBColor(24, 64, 88)
    run.font.size = Pt(15)
    info3.add_run(data['report_date'].strftime("%B %d, %Y"))
    
    info4 = make_tight(doc.add_paragraph())
    run = info4.add_run("Samples Taken:\n")
    run.font.name = 'Bebas Neue'
    run.bold = True
    run.font.color.rgb = RGBColor(24, 64, 88)
    run.font.size = Pt(15)
    info4.add_run("Exterior control sample (outdoor air)\n")
    for i, sample in enumerate(data['samples'], 1):
        info4.add_run(f"Sample {i}: {sample['type']} taken at {sample['location']}\n")
    
    # ===== PAGE 2: OFFICIAL LETTER =====
    doc.add_page_break()
    
    letter_header = doc.add_paragraph()
    run = letter_header.add_run("State Licensed Mold Assessment Consultant:\n")
    run.bold = True
    letter_header.add_run("Azeem Iqbal — TDLR MAC #2189\n\n")
    run = letter_header.add_run("Report Date:\n")
    run.bold = True
    letter_header.add_run(data['report_date'].strftime("%B %d, %Y"))
    
    doc.add_paragraph()
    doc.add_paragraph("To whom it may concern,")
    
    # Introduction paragraphs
    doc.add_paragraph(
        f"Mold Testing and Removal was hired to conduct a mold assessment at the property located at "
        f"{data['address']}, {data['city']}, {data['state']} {data['zip']}. The purpose of this assessment was to "
        f"evaluate the indoor air quality, identify potential sources of fungal growth, and provide recommendations for remediation."
    )
    
    doc.add_paragraph(
        "The assessment included a visual inspection, moisture mapping using a Protimeter Moisture Meter, "
        "and the collection of bioaerosol (air) and surface (swab) samples. Samples were collected from the "
        "interior of the property and the exterior for control purposes."
    )
    
    doc.add_paragraph("The samples were sent to PRO-LAB, an accredited laboratory, for viable mold/fungi analysis.")
    
    # Results paragraph - dynamic based on findings
    if lab_results['mold_present']:
        results_p = doc.add_paragraph()
        results_p.add_run("Based on the laboratory results and visual inspection, ")
        run = results_p.add_run("active mold growth was confirmed")
        run.bold = True
        run.italic = True
        results_p.add_run(" in the following areas:")
        
        for area in data['affected_areas']:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(f"{area['name']} — {area['finding']}")
        
        # Humidity warning
        if data['humidity'] > 50:
            humidity_p = doc.add_paragraph()
            run = humidity_p.add_run(
                f"The indoor relative humidity was recorded at {data['humidity']}%, which is above the recommended range "
                f"(30-50%) and conducive to microbial growth."
            )
            run.bold = True
        
        # Official notification
        doc.add_paragraph()
        notification = doc.add_paragraph()
        run = notification.add_run(
            "This letter serves as official notification that professional mold remediation is required to return "
            "the property to a normal fungal ecology (Condition 1). The property should be remediated by a State "
            "Licensed Mold Remediation Contractor (MRC) in accordance with the Texas Mold Assessment and Remediation Rules (TMARR)."
        )
        run.bold = True
    else:
        results_p = doc.add_paragraph()
        results_p.add_run("Based on the laboratory results and visual inspection, ")
        run = results_p.add_run("no significant mold contamination was identified")
        run.bold = True
        run.italic = True
        results_p.add_run(". The indoor spore counts are within acceptable levels compared to the outdoor control sample.")
        
        doc.add_paragraph(
            "This letter serves as confirmation that the property's fungal ecology is within normal parameters (Condition 1) "
            "at the time of inspection."
        )
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    sig = doc.add_paragraph()
    run = sig.add_run("Azeem Iqbal")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph("State of Texas Licensed Mold Assessment Consultant")
    doc.add_paragraph("TDLR MAC #2189 (Exp. 10/24/2027)")
    
    # ===== PAGE 3: VISUAL OBSERVATIONS =====
    doc.add_page_break()
    
    obs_title = make_tight(doc.add_paragraph())
    run = obs_title.add_run("Visual Observations & Moisture Readings")
    run.font.name = 'Bebas Neue'
    run.bold = True
    run.font.color.rgb = RGBColor(24, 64, 88)
    run.font.size = Pt(17)
    
    # Environmental conditions
    env_p = make_top_tight(doc.add_paragraph())
    run = env_p.add_run("Environmental Conditions: ")
    run.bold = True
    env_p.add_run(f"The indoor relative humidity (rH) was recorded at ")
    run = env_p.add_run(f"{data['humidity']}%")
    run.bold = True
    if data['humidity'] > 50:
        run.font.color.rgb = RGBColor(220, 53, 69)
    env_p.add_run(", which is ")
    if data['humidity'] > 50:
        env_p.add_run("above the recommended range (30-50%) and conducive to microbial growth.")
    else:
        env_p.add_run("within the recommended range (30-50%).")

    # Outdoor Control Sample placeholder
    ocs_title = make_tight(doc.add_paragraph())
    run = ocs_title.add_run("Outdoor Control Sample")
    run.font.name = 'Bebas Neue'
    run.bold = True
    run.font.color.rgb = RGBColor(24, 64, 88)
    run.font.size = Pt(17)
    ocs_p = make_top_tight(doc.add_paragraph())
    run = ocs_p.add_run("An air sample is taken outside to serve as a baseline for all other air samples to be compared against.")
    
    # Area observations with photos
    for area in data['affected_areas']:
        area_title = doc.add_paragraph()
        run = area_title.add_run(area['name'])
        run.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph(area['description'])
        
        # Add photos for this area if available
        area_key = area['name'].lower().replace(' ', '_')
        if photos.get(area_key):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(photos[area_key], width=Inches(3))
            except:
                pass
    
    # ===== PAGE 4: LAB RESULTS =====
    doc.add_page_break()
    
    lab_title = doc.add_paragraph()
    run = lab_title.add_run("Laboratory Results Analysis")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(24, 64, 88)
    
    doc.add_paragraph(
        "Samples were submitted to PRO-LAB (an accredited laboratory) for analysis. "
        "The following summarizes the findings compared to the outdoor control sample."
    )
    
    # Air Sample Table
    air_title = doc.add_paragraph()
    run = air_title.add_run("Air Sample Comparison (Bioaerosol)")
    run.bold = True
    run.font.size = Pt(12)
    
    air_table = doc.add_table(rows=1, cols=4)
    air_table.style = 'Table Grid'
    
    # Header row
    header_cells = air_table.rows[0].cells
    headers = ['Location', 'Fungal Type', 'Spores/m³', 'Interpretation']
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], "D5E8F0")
    
    # Add sample data rows
    for sample in data['air_samples']:
        row = air_table.add_row()
        row.cells[0].text = sample['location']
        row.cells[1].text = sample['fungal_type']
        row.cells[2].text = str(sample['spore_count'])
        row.cells[3].text = sample['interpretation']
        
        if sample['interpretation'] == 'ELEVATED':
            set_cell_shading(row.cells[2], "FFCCCC")
            set_cell_shading(row.cells[3], "FFCCCC")
            row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
            row.cells[3].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Surface Sample Table
    surface_title = doc.add_paragraph()
    run = surface_title.add_run("Surface Sample Results (Swab)")
    run.bold = True
    run.font.size = Pt(12)
    
    surface_table = doc.add_table(rows=1, cols=3)
    surface_table.style = 'Table Grid'
    
    header_cells = surface_table.rows[0].cells
    headers = ['Location', 'Sample Type', 'Result']
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], "D5E8F0")
    
    for sample in data['surface_samples']:
        row = surface_table.add_row()
        row.cells[0].text = sample['location']
        row.cells[1].text = 'Swab'
        row.cells[2].text = sample['result']
        
        if 'UNUSUAL' in sample['result'] or 'Mold Present' in sample['result']:
            set_cell_shading(row.cells[2], "FFCCCC")
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
            row.cells[2].paragraphs[0].runs[0].bold = True
    
    # Mold Types section
    doc.add_paragraph()
    mold_title = doc.add_paragraph()
    run = mold_title.add_run("Mold Types Identified")
    run.bold = True
    run.font.size = Pt(12)
    
    for mold_type in data.get('mold_types_found', []):
        mold_p = doc.add_paragraph()
        run = mold_p.add_run(f"{mold_type['name']}: ")
        run.bold = True
        if mold_type.get('dangerous'):
            run.font.color.rgb = RGBColor(220, 53, 69)
        mold_p.add_run(mold_type['description'])
    
    # ===== PAGE 5: CONCLUSIONS & RECOMMENDATIONS =====
    doc.add_page_break()
    
    conc_title = doc.add_paragraph()
    run = conc_title.add_run("Conclusions")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(24, 64, 88)
    
    doc.add_paragraph("Based on the visual inspection, moisture readings, and laboratory results, the following conclusions are made:")
    
    for i, conclusion in enumerate(data['conclusions'], 1):
        conc_p = doc.add_paragraph()
        run = conc_p.add_run(f"{i}. {conclusion['area']}: ")
        run.bold = True
        conc_p.add_run(conclusion['finding'])
    
    doc.add_paragraph()
    
    rec_title = doc.add_paragraph()
    run = rec_title.add_run("Recommendations")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(24, 64, 88)
    
    if lab_results['mold_present']:
        doc.add_paragraph("To return the property to a normal fungal ecology (Condition 1), the following remediation steps are recommended:")
        
        recommendations = [
            ("Professional Remediation", "Hire a State Licensed Mold Remediation Contractor (MRC) to prepare a work plan based on this protocol."),
            ("Containment", "Establish critical barriers (polyethylene sheeting) around affected areas to prevent spore dispersion. Establish negative air pressure."),
            ("Removal", "Remove and discard affected drywall and materials. Continue removal 2 feet beyond visible growth."),
            ("Cleaning", "HEPA vacuum and damp-wipe all remaining structural surfaces within the containment."),
            ("Humidity Control", "Dehumidification is required to lower the indoor RH to between 30-50%."),
            ("Clearance Testing", "After remediation, a Post-Remediation Assessment (clearance test) must be performed by a TDLR Mold Assessment Consultant.")
        ]
    else:
        doc.add_paragraph("Based on the findings, the following recommendations are made:")
        recommendations = [
            ("Humidity Control", "Maintain indoor relative humidity between 30-50% to prevent future mold growth."),
            ("Regular Inspection", "Periodically check areas prone to moisture for signs of water intrusion or condensation."),
            ("Ventilation", "Ensure proper ventilation in bathrooms, kitchens, and laundry areas.")
        ]
    
    for rec_title_text, rec_desc in recommendations:
        rec_p = doc.add_paragraph(style='List Bullet')
        run = rec_p.add_run(f"{rec_title_text}: ")
        run.bold = True
        rec_p.add_run(rec_desc)
    
    # Compliance note
    doc.add_paragraph()
    compliance = doc.add_paragraph()
    run = compliance.add_run(
        "This report is generated in accordance with the Texas Mold Assessment and Remediation Rules (TMARR). "
        "Limitations: This inspection is limited to the areas accessible at the time of inspection."
    )
    run.italic = True
    run.font.size = Pt(9)
    
    # ===== PAGE 6: TERMS AND CONDITIONS =====
    doc.add_page_break()
    
    terms_title = doc.add_paragraph()
    run = terms_title.add_run("Terms and Conditions")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(24, 64, 88)
    
    terms = [
        ("Inspection Limitation", "This inspection and the information set forth in the report is provided solely for the purpose of verifying that certain structural or physical characteristics exist at the Location Address listed. The undersigned and company representative does not make a health or safety certification or warranty, express or implied, of any kind."),
        ("Limitation of Liability", "The Client agrees that Inspector's liability for errors and/or omissions shall be limited to the maximum of a full refund of the fee paid for the inspection. The Client agrees to assume all risk of loss which exceeds the fee paid."),
        ("Sampling Limitations", "Mold spore sampling results represent conditions at the time and location of sampling only. Conditions can change rapidly due to environmental factors, occupant activities, and remediation efforts."),
        ("Health Disclaimer", "This report does not constitute medical advice. Individuals with health concerns related to potential mold exposure should consult with a qualified healthcare professional."),
        ("Report Usage", "This report is prepared exclusively for the named client and may not be reproduced or distributed to third parties without written consent from Mold Testing and Removal.")
    ]
    
    for term_title, term_text in terms:
        term_p = doc.add_paragraph()
        run = term_p.add_run(f"{term_title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = term_p.add_run(term_text)
        run2.font.size = Pt(10)
    
    # Lab report reference
    doc.add_paragraph()
    lab_ref = doc.add_paragraph()
    lab_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lab_ref.add_run("— Laboratory Report Attached —")
    run.bold = True
    run.italic = True
    
    lab_ref2 = doc.add_paragraph()
    lab_ref2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lab_ref2.add_run("PRO-LAB Certificate of Mold Analysis follows this page")
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Mold Testing and Removal\n")
    run.bold = True
    run.font.color.rgb = RGBColor(24, 64, 88)
    footer.add_run("2031 John West Rd, Suite 119 | Dallas, TX 75228\n")
    footer.add_run("(817) 718-5086 | Azeem@RestorationCleanupService.com")
    
    # Save Word doc to bytes
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer

def merge_pdfs(docx_buffer, lab_pdf_bytes):
    """Convert Word doc to PDF and merge with lab report"""
    # For now, we'll return the Word doc and lab PDF separately
    # Full PDF merging would require additional dependencies
    return docx_buffer, lab_pdf_bytes

# ===== MAIN APP =====

# Tabs for different sections
tab1, tab2, tab3, tab4 = st.tabs(["📋 Client Info", "📸 Photos & Lab Report", "🔬 Lab Results", "📄 Generate Report"])

with tab1:
    st.markdown('<p class="section-header">Client & Property Information</p>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        client_name = st.text_input("Client Name *", placeholder="e.g., Tarah Dickerson")
        address = st.text_input("Property Address *", placeholder="e.g., 2073 Bowie St")
        city = st.text_input("City *", placeholder="e.g., Sanger")
        
    with col2:
        phone = st.text_input("Phone", placeholder="(214) 886-2801")
        email = st.text_input("Email", placeholder="client@email.com")
        state = st.selectbox("State", ["TX", "OK", "AR", "LA", "NM"], index=0)
        zip_code = st.text_input("ZIP Code *", placeholder="76266")
    
    st.markdown('<p class="section-header">Inspection Details</p>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        inspection_date = st.date_input("Inspection Date *", value=date.today())
    with col2:
        report_date = st.date_input("Report Date *", value=date.today())
    with col3:
        humidity = st.number_input("Indoor Humidity (%)", min_value=0, max_value=100, value=62)
    
    st.markdown('<p class="section-header">Sample Locations</p>', unsafe_allow_html=True)
    
    num_samples = st.number_input("Number of samples taken (excluding outdoor control)", min_value=1, max_value=10, value=3)
    
    samples = []
    for i in range(int(num_samples)):
        col1, col2 = st.columns(2)
        with col1:
            sample_type = st.selectbox(f"Sample {i+1} Type", ["Air Sample", "Swab"], key=f"type_{i}")
        with col2:
            sample_location = st.text_input(f"Sample {i+1} Location", placeholder=f"e.g., Kitchen Sink Cabinet", key=f"loc_{i}")
        samples.append({"type": sample_type, "location": sample_location})

with tab2:
    st.markdown('<p class="section-header">Upload Photos</p>', unsafe_allow_html=True)
    
    property_photo = st.file_uploader("Property Photo (exterior)", type=['jpg', 'jpeg', 'png'], key="property")
    
    st.markdown("**Inspection Photos**")
    
    inspection_photos = {}
    num_areas = st.number_input("Number of affected areas to document", min_value=1, max_value=10, value=3)
    
    for i in range(int(num_areas)):
        col1, col2, col3 = st.columns([2, 2, 3])
        with col1:
            area_name = st.text_input(f"Area {i+1} Name", placeholder="e.g., Kitchen", key=f"area_name_{i}")
        with col2:
            area_finding = st.selectbox(f"Area {i+1} Finding", 
                ["Active mold growth confirmed", "Elevated spore counts", "Visual mold present", "No mold detected"],
                key=f"area_finding_{i}")
        with col3:
            area_photo = st.file_uploader(f"Photo for Area {i+1}", type=['jpg', 'jpeg', 'png'], key=f"area_photo_{i}")
        
        area_desc = st.text_area(f"Area {i+1} Description", 
            placeholder="Describe observations, moisture readings, etc.", 
            key=f"area_desc_{i}", height=80)
        
        if area_name:
            inspection_photos[f"area_{i}"] = {
                "name": area_name,
                "finding": area_finding,
                "photo": area_photo,
                "description": area_desc
            }
        
        st.divider()
    
    st.markdown('<p class="section-header">Upload PRO-LAB Report</p>', unsafe_allow_html=True)
    
    lab_pdf = st.file_uploader("PRO-LAB Certificate of Mold Analysis (PDF) *", type=['pdf'], key="lab_pdf")
    
    if lab_pdf:
        st.success("✅ Lab report uploaded successfully!")
        
        if st.button("🔍 Analyze Lab Report"):
            with st.spinner("Analyzing lab report..."):
                results = parse_prolab_pdf(lab_pdf)
                st.session_state.lab_results = results
                
                if results['mold_present']:
                    st.error("⚠️ MOLD DETECTED - Remediation Required")
                else:
                    st.success("✅ No significant mold contamination detected")
                
                if results['dangerous_mold']:
                    st.error("🚨 DANGEROUS MOLD (Stachybotrys/Chaetomium) DETECTED!")

with tab3:
    st.markdown('<p class="section-header">Lab Results Entry</p>', unsafe_allow_html=True)
    
    if st.session_state.lab_results:
        st.info("Lab report has been analyzed. You can adjust values below if needed.")
    
    st.markdown("**Air Sample Results**")
    
    # Outdoor control
    col1, col2, col3 = st.columns(3)
    with col1:
        outdoor_pen_asp = st.number_input("Outdoor Control - Pen/Asp (spores/m³)", value=160, key="outdoor_pen")
    with col2:
        outdoor_total = st.number_input("Outdoor Control - Total Spores", value=266, key="outdoor_total")
    with col3:
        outdoor_interpretation = st.selectbox("Outdoor Interpretation", ["Baseline (Reference)"], key="outdoor_int")
    
    st.markdown("**Indoor Air Samples**")
    
    air_samples = [{"location": "Outdoor Control", "fungal_type": "Penicillium/Aspergillus", 
                    "spore_count": outdoor_pen_asp, "interpretation": "Baseline (Reference)"}]
    
    num_air_samples = st.number_input("Number of indoor air samples", min_value=1, max_value=5, value=1)
    
    for i in range(int(num_air_samples)):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            loc = st.text_input(f"Indoor Air {i+1} Location", value="Closet Air Sample", key=f"air_loc_{i}")
        with col2:
            fungal = st.selectbox(f"Fungal Type", ["Penicillium/Aspergillus", "Cladosporium", "Other Ascospores", "Stachybotrys", "Chaetomium"], key=f"air_fungal_{i}")
        with col3:
            count = st.number_input(f"Spores/m³", value=210, key=f"air_count_{i}")
        with col4:
            interp = st.selectbox(f"Interpretation", ["ELEVATED", "Not Elevated", "Baseline"], key=f"air_interp_{i}")
        
        air_samples.append({"location": loc, "fungal_type": fungal, "spore_count": count, "interpretation": interp})
    
    st.markdown("**Surface Sample Results (Swab)**")
    
    surface_samples = []
    num_surface = st.number_input("Number of surface samples", min_value=0, max_value=5, value=2)
    
    for i in range(int(num_surface)):
        col1, col2 = st.columns(2)
        with col1:
            surf_loc = st.text_input(f"Surface {i+1} Location", key=f"surf_loc_{i}")
        with col2:
            surf_result = st.selectbox(f"Result", ["UNUSUAL / Mold Present", "UNUSUAL / Mold Present (Stachybotrys)", "Normal"], key=f"surf_result_{i}")
        
        if surf_loc:
            surface_samples.append({"location": surf_loc, "result": surf_result})
    
    st.markdown("**Mold Types Found**")
    
    mold_options = st.multiselect("Select mold types identified in the report",
        ["Penicillium/Aspergillus", "Cladosporium", "Chaetomium", "Stachybotrys", "Other Ascospores"],
        default=["Penicillium/Aspergillus"])
    
    mold_descriptions = {
        "Penicillium/Aspergillus": ("The most common mold species in indoor air samples. Often associated with water damage and elevated humidity. Known allergen (Type I and Type III).", False),
        "Cladosporium": ("The most common spore type worldwide. Commonly found on wood and wallboard. Known allergen but also common outdoors.", False),
        "Chaetomium": ("A water-indicating mold found on cellulose materials. Should not be observed indoors unless building materials have been wetted.", True),
        "Stachybotrys": ("Known as 'black mold.' Requires high water content to grow. A water-indicating mold that produces mycotoxins. Professional remediation required.", True),
        "Other Ascospores": ("Spores from a large group of fungi common everywhere outdoors. When found indoors in higher concentrations, they may indicate moisture issues.", False)
    }

with tab4:
    st.markdown('<p class="section-header">Generate Final Report</p>', unsafe_allow_html=True)
    
    # Validation
    missing_fields = []
    if not client_name:
        missing_fields.append("Client Name")
    if not address:
        missing_fields.append("Address")
    if not city:
        missing_fields.append("City")
    if not zip_code:
        missing_fields.append("ZIP Code")
    if not lab_pdf:
        missing_fields.append("PRO-LAB Report PDF")
    
    if missing_fields:
        st.warning(f"⚠️ Please fill in the following required fields: {', '.join(missing_fields)}")
    
    # Summary preview
    st.markdown("### Report Preview")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Client Information**")
        st.write(f"Name: {client_name or 'Not entered'}")
        st.write(f"Address: {address or 'Not entered'}, {city or ''}, {state} {zip_code or ''}")
        st.write(f"Inspection Date: {inspection_date}")
        st.write(f"Humidity: {humidity}%")
    
    with col2:
        st.markdown("**Lab Results Summary**")
        if st.session_state.lab_results:
            if st.session_state.lab_results['mold_present']:
                st.error("⚠️ Mold Detected - Remediation Required")
            else:
                st.success("✅ No Significant Mold")
            
            if st.session_state.lab_results['dangerous_mold']:
                st.error("🚨 Dangerous Mold Present!")
        else:
            st.info("Upload and analyze lab report to see results")
    
    # Conclusions input
    st.markdown("### Conclusions")
    st.info("Enter conclusions for each affected area:")
    
    conclusions = []
    for i, (key, area_data) in enumerate(inspection_photos.items()):
        if area_data.get('name'):
            conclusion_text = st.text_area(
                f"Conclusion for {area_data['name']}", 
                value=f"{area_data.get('finding', '')}. {area_data.get('description', '')}",
                key=f"conclusion_{i}",
                height=80
            )
            conclusions.append({"area": area_data['name'], "finding": conclusion_text})
    
    # Add humidity conclusion
    if humidity > 50:
        conclusions.append({
            "area": "Humidity",
            "finding": f"The indoor relative humidity of {humidity}% is too high and is contributing to fungal proliferation."
        })
    
    st.markdown("---")
    
    # Generate button
    if st.button("🚀 Generate Report", type="primary", disabled=len(missing_fields) > 0):
        with st.spinner("Generating your report..."):
            
            # Prepare data
            data = {
                'client_name': client_name,
                'address': address,
                'city': city,
                'state': state,
                'zip': zip_code,
                'phone': phone,
                'email': email,
                'inspection_date': inspection_date,
                'report_date': report_date,
                'humidity': humidity,
                'samples': samples,
                'affected_areas': [
                    {
                        'name': area_data['name'],
                        'finding': area_data['finding'],
                        'description': area_data.get('description', '')
                    }
                    for area_data in inspection_photos.values() if area_data.get('name')
                ],
                'air_samples': air_samples,
                'surface_samples': surface_samples,
                'mold_types_found': [
                    {'name': m, 'description': mold_descriptions[m][0], 'dangerous': mold_descriptions[m][1]}
                    for m in mold_options
                ],
                'conclusions': conclusions
            }
            
            # Prepare photos
            photos = {}
            if property_photo:
                photos['property'] = property_photo
            
            for key, area_data in inspection_photos.items():
                if area_data.get('photo'):
                    photos[key] = area_data['photo']
            
            # Get lab results
            lab_results = st.session_state.lab_results or {
                'mold_present': any(s.get('interpretation') == 'ELEVATED' for s in air_samples) or
                               any('UNUSUAL' in s.get('result', '') for s in surface_samples),
                'dangerous_mold': 'Stachybotrys' in mold_options or 'Chaetomium' in mold_options
            }
            
            # Read lab PDF bytes
            lab_pdf_bytes = lab_pdf.read() if lab_pdf else None
            lab_pdf.seek(0) if lab_pdf else None
            
            # Generate report
            try:
                docx_buffer = create_report(data, photos, lab_pdf_bytes, lab_results)
                
                st.success("✅ Report generated successfully!")
                
                # Download buttons
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        label="📥 Download Word Report (.docx)",
                        data=docx_buffer.getvalue(),
                        file_name=f"{client_name.replace(' ', '_')}_Mold_Assessment_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                with col2:
                    if lab_pdf_bytes:
                        st.download_button(
                            label="📥 Download Lab Report (.pdf)",
                            data=lab_pdf_bytes,
                            file_name=f"{client_name.replace(' ', '_')}_Lab_Report.pdf",
                            mime="application/pdf"
                        )
                
                st.info("💡 **Tip:** Open the Word document, review it, then combine it with the Lab Report PDF using Adobe Acrobat or a free PDF merger tool.")
                
            except Exception as e:
                st.error(f"Error generating report: {str(e)}")
                st.exception(e)

# Sidebar with instructions
with st.sidebar:
    st.markdown("## 📖 Instructions")
    st.markdown("""
    1. **Client Info Tab**: Enter client and property details
    2. **Photos Tab**: Upload property photo, inspection photos, and the PRO-LAB PDF
    3. **Lab Results Tab**: Review/adjust the extracted lab data
    4. **Generate Tab**: Review summary and generate the report
    
    ---
    
    ## ✅ Required Fields
    - Client Name
    - Property Address
    - City & ZIP
    - PRO-LAB PDF Report
    
    ---
    
    ## 🔬 Lab Analysis
    The app will automatically:
    - Detect if mold is present
    - Identify dangerous molds (Stachybotrys, Chaetomium)
    - Determine if levels are ELEVATED
    - Adjust conclusions accordingly
    
    ---
    
    ## 📞 Support
    Mold Testing and Removal
    (817) 718-5086
    """)
